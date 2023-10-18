from PyQt5.QtWidgets import QMainWindow,QApplication, QTableWidgetItem, QHeaderView, QTableWidget, QMessageBox, QAbstractItemView, QComboBox, QDoubleSpinBox, QTextEdit, QRadioButton, QTableView
from PyQt5.QtCore import Qt, QTimer
from PyQt5.QtGui import QIcon
from PyQt5 import QtCore, QtWidgets
import pandas as pd
import json
import re
import docx
import subprocess
import shutil
from lxml import etree
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Pt, Inches
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from PyQt5 import QtWidgets
import sys, os ,webbrowser ,subprocess

def guardar_valores_en_json(windows, archivo_json, rango_le, rango_te, rango_rb, rango_tie, rango_de, tabla_datos, tabla_ofertas):
	datos = {}

	# Guardar valores de QLineEdit
	for i in range(*rango_le):
		nombre_widget = f"Le{i}"
		widget = windows.findChild(QtWidgets.QLineEdit, nombre_widget)
		valor = widget.text()
		datos[nombre_widget] = valor

	# Guardar valores de QTextEdit
	for i in range(*rango_te):
		nombre_widget = f"Te{i}"
		widget = windows.findChild(QtWidgets.QTextEdit, nombre_widget)
		valor = widget.toPlainText()
		datos[nombre_widget] = valor

	# Guardar valores de QRadioButton
	for i in range(*rango_rb):
		nombre_widget = f"Rb{i}"
		widget = windows.findChild(QtWidgets.QRadioButton, nombre_widget)
		valor = widget.isChecked()
		datos[nombre_widget] = valor

	# Guardar valores de QTimeEdit
	for i in range(*rango_tie):
		nombre_widget = f"Tie{i}"
		widget = windows.findChild(QtWidgets.QTimeEdit, nombre_widget)
		valor = widget.time().toString()
		datos[nombre_widget] = valor

	# Guardar valores de QDateEdit
	for i in range(*rango_de):
		nombre_widget = f"De{i}"
		widget = windows.findChild(QtWidgets.QDateEdit, nombre_widget)
		valor = widget.date().toString("dd-MM-yyyy")
		datos[nombre_widget] = valor

	# Obtener los valores de la tabla y agregarlos al diccionario
	for fila in range(tabla_datos.rowCount()):
		for columna in range(tabla_datos.columnCount()):
			clave = f"tabla{fila}.{columna}"
			item = tabla_datos.item(fila, columna)
			valor = item.text() if item is not None else ""
			datos[clave] = valor

	columna = 1  # Columna a ordenar
	valores_columna = []
	for fila in range(tabla_ofertas.rowCount()):
		item = tabla_ofertas.item(fila, columna)
		if item is not None and item.text():
			valor = item.text().replace(",", ".")  # Reemplazar comas por puntos
			valores_columna.append(valor)
		else:
			valores_columna.append('')  # Agregar espacio vacío para valores vacíos

	# Obtener el orden clasificatorio de los valores no vacíos
	orden_clasificatorio = sorted(range(len(valores_columna)), key=lambda k: float(valores_columna[k]) if valores_columna[k] else float('inf'))

	# Almacenar los valores junto con su orden de clasificación en un vector
	vector_ordenado = ['' for _ in range(tabla_ofertas.rowCount())]  # Inicializar vector_ordenado con valores vacíos
	contador = 1
	for orden in orden_clasificatorio:
		if valores_columna[orden]:  # Si el valor no está vacío
			vector_ordenado[orden] = str(contador)
			contador += 1


	# Obtener los valores de la tabla y agregarlos al diccionario
	for fila in range(tabla_ofertas.rowCount()):
		item = tabla_ofertas.item(fila, 1)
		if item is not None:
			valor = item.text()
		else:
			valor = ""

		clave = f"tabla{fila}.5"
		datos[clave] = valor

		clave_orden = f"tabla{fila}.6"
		datos[clave_orden] = vector_ordenado[fila]

		clave_oferta = f"tabla{fila}.7"
		datos[clave_oferta] = 'Sí' if valor else 'No'

	with open(archivo_json, "w") as archivo:
		json.dump(datos, archivo)

def cargar_valores_desde_json( windows,archivo_json, rango_le, rango_te, rango_rb, rango_tie, rango_de, tabla_datos, tabla_ofertas):
	try:
		if os.path.exists(archivo_json):
			with open(archivo_json, "r") as archivo:
				datos = json.load(archivo)

			windows.lista_filas = datos.get("lista_filas", [])
			windows.datos_json = datos  # Guardar los datos del archivo JSON

			# Cargar valores de QLineEdit
			for i in range(*rango_le):
				nombre_widget = f"Le{i}"
				widget = windows.findChild(QtWidgets.QLineEdit, nombre_widget)
				valor = datos.get(nombre_widget)

				if valor is not None:
					widget.setText(valor)

			# Cargar valores de QTextEdit
			for i in range(*rango_te):
				nombre_widget = f"Te{i}"
				widget = windows.findChild(QtWidgets.QTextEdit, nombre_widget)
				valor = datos.get(nombre_widget)
				if valor is not None:
					widget.setPlainText(valor)

			# Cargar valores de QRadioButton
			for i in range(*rango_rb):
				nombre_widget = f"Rb{i}"
				widget =windows.findChild(QtWidgets.QRadioButton, nombre_widget)
				valor = datos.get(nombre_widget)
				if valor is not None:
					widget.setChecked(valor)

			# Cargar valores de QTimeEdit
			for i in range(*rango_tie):
				nombre_widget = f"Tie{i}"
				widget = windows.findChild(QtWidgets.QTimeEdit, nombre_widget)
				valor = datos.get(nombre_widget)
				if valor is not None:
					tiempo = QtCore.QTime.fromString(valor)
					widget.setTime(tiempo)

			# Cargar valores de QDateEdit
			for i in range(*rango_de):
				nombre_widget = f"De{i}"
				widget = windows.findChild(QtWidgets.QDateEdit, nombre_widget)
				valor = datos.get(nombre_widget)
				if valor is not None:
					fecha = QtCore.QDate.fromString(valor, "dd-MM-yyyy")
					widget.setDate(fecha)

			# Cargar valores de la tabla desde el diccionario
			for fila in range(tabla_datos.rowCount()):
				for columna in range(tabla_datos.columnCount()):
					clave = f"tabla{fila}.{columna}"
					valor = datos.get(clave)

					if valor is not None:
						item = QtWidgets.QTableWidgetItem(valor)
						tabla_datos.setItem(fila, columna, item)

			# Cargar valores en la segunda columna de tabla_ofertas
			for fila in range(tabla_ofertas.rowCount()):
				clave = f"tabla{fila}.5"
				valor = datos.get(clave)

				if valor is not None:
					item = QtWidgets.QTableWidgetItem(valor)
					tabla_ofertas.setItem(fila, 1, item)

		else:
			print(f"No se encontró el archivo JSON: {archivo_json}")

	except Exception as e:
		QtWidgets.QMessageBox.critical(None, "Error", f"Ocurrió un error al cargar valores desde JSON: {str(e)}")
		
def Sustituir_cartas(self, archivo_entrada, archivo_salida, archivoJson):
		
		try:
			# Leer el archivo JSON
			with open(archivoJson) as f:
				datos_json = json.load(f)
		except FileNotFoundError:
			print(f"Archivo JSON no encontrado: {archivoJson}")
			return
		except json.JSONDecodeError:
			print(f"Error al decodificar el archivo JSON: {archivoJson}")
			return
		for i_str  in range(1, self.TwEmpresas.rowCount() +1 ):
		#for i in ["0", "1", "2", "3", "4"]:
			doc = Document(archivo_entrada)
			i = str(i_str-1)
			# Recorrer todos los párrafos del documento y buscar y reemplazar cada clave del archivo JSON con su valor correspondiente
			for paragraph in doc.paragraphs:
				if paragraph.runs:
					# Guardar el estilo actual de fuente y tamaño
					font_name = paragraph.runs[0].font.name
					font_size = paragraph.runs[0].font.size

					for buscar, reemplazar in datos_json.items():
						if buscar in ["tabla"+i+".0", "tabla"+i+".1", "tabla"+i+".2", "tabla"+i+".3"]:
							buscar_str = "tabla" + str(int(buscar.split(".")[1]) )
							reemplazar_str = str(reemplazar)
							regex = re.compile(r'\b' + re.escape(buscar_str) + r'\b')
							paragraph.text = regex.sub(reemplazar_str, paragraph.text)

					# Restaurar el estilo de fuente y tamaño
					for run in paragraph.runs:
						run.font.name = font_name
						run.font.size = font_size

			# Guardar el documento modificado en un archivo de salida
			doc.save('{}{}{}'.format(self.path+'\\1_Salida\\', i, archivo_salida))
			os.startfile('{}{}{}'.format(self.path+'\\1_Salida\\', i, archivo_salida))
			
def Sustituir_cartas_adj_no(self, archivo_entrada_no,archivo_entrada_adj,empresa_adj, archivo_salida_no, archivo_salida_adj, archivoJson):
		print(empresa_adj)
		try:
			# Leer el archivo JSON
			with open(archivoJson) as f:
				datos_json = json.load(f)
		except FileNotFoundError:
			print(f"Archivo JSON no encontrado: {archivoJson}")
			return
		except json.JSONDecodeError:
			print(f"Error al decodificar el archivo JSON: {archivoJson}")
			return
		#for i in ["0", "1", "2", "3", "4"]:
		for i_str  in range(1, self.TwEmpresas.rowCount() +1 ):
			print(str(i_str)+" y "+str(empresa_adj))
			print(str(empresa_adj) + " y " + str(i_str))
			if i_str ==empresa_adj:
			
				doc = Document(archivo_entrada_adj)
			else :
				doc = Document(archivo_entrada_no)
			i = str(i_str-1)
			# Recorrer todos los párrafos del documento y buscar y reemplazar cada clave del archivo JSON con su valor correspondiente
			for paragraph in doc.paragraphs:
				if paragraph.runs:
					# Guardar el estilo actual de fuente y tamaño
					font_name = paragraph.runs[0].font.name
					font_size = paragraph.runs[0].font.size

					for buscar, reemplazar in datos_json.items():
						if buscar in ["tabla"+i+".0", "tabla"+i+".1", "tabla"+i+".2", "tabla"+i+".3"]:
							buscar_str = "tabla" + str(int(buscar.split(".")[1]) )
							reemplazar_str = str(reemplazar)
							regex = re.compile(r'\b' + re.escape(buscar_str) + r'\b')
							paragraph.text = regex.sub(reemplazar_str, paragraph.text)

					# Restaurar el estilo de fuente y tamaño
					for run in paragraph.runs:
						run.font.name = font_name
						run.font.size = font_size
			if i_str ==empresa_adj:
				print("adj")
				doc.save('{}{}{}'.format(self.path+'\\1_Salida\\', i, archivo_salida_adj))
				os.startfile('{}{}{}'.format(self.path+'\\1_Salida\\', i, archivo_salida_adj))
			else :
				print("no_adj")
				doc.save('{}{}{}'.format(self.path+'\\1_Salida\\', i, archivo_salida_no))
				os.startfile('{}{}{}'.format(self.path+'\\1_Salida\\', i, archivo_salida_no))
			
def Sustituir(self, ruta_documento_origen, ruta_documento_destino, archivo_json):
	try:
		doc = docx.Document(ruta_documento_origen)

		# Cargar los datos desde el archivo JSON
		with open(archivo_json, 'r') as json_file:
			datos_json = json.load(json_file)
		# Recorrer las tablas en el documento
		for table in doc.tables:
			for row in table.rows:
				for cell in row.cells:
					for paragraph in cell.paragraphs:
						if paragraph.runs:
							# Guardar el estilo actual de fuente y tamaño
							font_name = paragraph.runs[0].font.name
							font_size = paragraph.runs[0].font.size

							# Buscar y reemplazar en el texto del párrafo utilizando expresiones regulares
							for buscar, reemplazar in datos_json.items():
								buscar_str = str(buscar)
								reemplazar_str = str(reemplazar)
								
								if es_cadena_numerica(reemplazar_str) and reemplazar_str.strip():#formatea los numeros
									valor_numerico = float(reemplazar_str)
									if valor_numerico > 100:
										reemplazar_str = formatear_numero(float(reemplazar_str))

								regex = re.compile(r'\b' + re.escape(buscar_str) + r'\b')
								paragraph.text = regex.sub(reemplazar_str, paragraph.text)

							# Restaurar el estilo de fuente y tamaño
							for run in paragraph.runs:
								run.font.name = font_name
								run.font.size = font_size

		for paragraph in doc.paragraphs:
			if paragraph.runs:
				# Guardar el estilo actual de fuente y tamaño
				font_name = paragraph.runs[0].font.name
				font_size = paragraph.runs[0].font.size

				# Buscar y reemplazar en el texto del párrafo utilizando expresiones regulares
				for buscar, reemplazar in datos_json.items():
					
						
					
					
					buscar_str = str(buscar)
					reemplazar_str = str(reemplazar)


					regex = re.compile(r'\b' + re.escape(buscar_str) + r'\b')
					paragraph.text = regex.sub(reemplazar_str, paragraph.text)

				# Restaurar el estilo de fuente y tamaño
				for run in paragraph.runs:
					run.font.name = font_name
					run.font.size = font_size

		doc.save(ruta_documento_destino)

		directorio = os.path.dirname(ruta_documento_destino)
		comando = f'start "" /D "{directorio}" "{ruta_documento_destino}"'
		subprocess.run(comando, shell=True)
	except Exception as e:
		error_message = f"Error al ejecutar el reemplazo: {e}"
		print(error_message)
		
def Sustituir_temp(self, ruta_documento_origen, ruta_documento_destino, archivo_json):
			try:
				doc = docx.Document(ruta_documento_origen)

				# Cargar los datos desde el archivo JSON
				with open(archivo_json, 'r') as json_file:
					datos_json = json.load(json_file)
				# Recorrer las tablas en el documento
				for table in doc.tables:
					for row in table.rows:
						for cell in row.cells:
							for paragraph in cell.paragraphs:
								if paragraph.runs:
									# Guardar el estilo actual de fuente y tamaño
									font_name = paragraph.runs[0].font.name
									font_size = paragraph.runs[0].font.size

									# Buscar y reemplazar en el texto del párrafo utilizando expresiones regulares
									for buscar, reemplazar in datos_json.items():
										buscar_str = str(buscar)
										reemplazar_str = str(reemplazar)
										if es_cadena_numerica(reemplazar_str) and reemplazar_str.strip():#formatea los numeros
											valor_numerico = float(reemplazar_str)
											if valor_numerico > 100:
												reemplazar_str = formatear_numero(float(reemplazar_str))
										regex = re.compile(r'\b' + re.escape(buscar_str) + r'\b')
										paragraph.text = regex.sub(reemplazar_str, paragraph.text)

									# Restaurar el estilo de fuente y tamaño
									for run in paragraph.runs:
										run.font.name = font_name
										run.font.size = font_size

				for paragraph in doc.paragraphs:
					if paragraph.runs:
						# Guardar el estilo actual de fuente y tamaño
						font_name = paragraph.runs[0].font.name
						font_size = paragraph.runs[0].font.size

						# Buscar y reemplazar en el texto del párrafo utilizando expresiones regulares
						for buscar, reemplazar in datos_json.items():
							buscar_str = str(buscar)
							reemplazar_str = str(reemplazar)
							regex = re.compile(r'\b' + re.escape(buscar_str) + r'\b')
							paragraph.text = regex.sub(reemplazar_str, paragraph.text)

						# Restaurar el estilo de fuente y tamaño
						for run in paragraph.runs:
							run.font.name = font_name
							run.font.size = font_size

				doc.save(ruta_documento_destino)

			except Exception as e:
				error_message = f"Error al ejecutar el reemplazo: {e}"
				print(error_message)
				
def es_numero_entero(valor):
    try:
        float_valor = float(valor)  # Intenta convertir la cadena a un número de punto flotante
        return float_valor.is_integer()  # Comprueba si el número es entero
    except ValueError:
        return False  # Si hay una excepción, el valor no es un número
def es_cadena_numerica(cadena):
    return all(char.isdigit() or char == '.' or char.isspace() for char in cadena)


def formatear_numero(numero):
    partes = str(numero).split('.')
    
    # Formatea la parte entera con puntos como separadores de miles
    parte_entera = '{:,}'.format(int(partes[0])).replace(',', '.')

    # Une la parte entera y la parte decimal con una coma
    if len(partes) > 1:
        resultado = parte_entera + ',' + partes[1]
    else:
        resultado = parte_entera

    return resultado

from PyQt5.QtGui import QDoubleValidator, QIntValidator
from PyQt5.QtWidgets import QHeaderView ,QTableWidgetItem, QTableWidget, QApplication
from PyQt5 import uic
from PyQt5.QtGui import QIcon, QPixmap
from PyQt5.QtWidgets import QMessageBox
from PyQt5.QtCore import Qt
from PyQt5.QtWidgets import QStyledItemDelegate
from PyQt5.QtGui import QIcon, QPixmap, QDoubleValidator, QIntValidator ,QValidator

from docx.shared import Inches

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Pt
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx import Document
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml.etree import tostring
from docx.shared import Pt
import PyQt5.QtWidgets as QtWidgets

def insertar_tabla_en_documento_inicio(archivo_entrada, archivo_salida, contenido_objetivo, filas):

	try:
		# Abrir el archivo de entrada
		documento = Document(archivo_entrada)

		# Crear una tabla con filas+1 y 4 columnas
		columnas = 4
		tabla = documento.add_table(rows=filas+1, cols=columnas)

		# Establecer el ancho de columna
		ancho_columna = Inches(2)
		for columna in tabla.columns:
			columna.width = round(ancho_columna)

		# Aplicar un borde a todas las celdas de la tabla
		for fila in tabla.rows:
			for celda in fila.cells:
				celda_borde = celda._element
				celda_borde.set(qn('w:borders'), tostring(OxmlElement('w:borders'), encoding='unicode'))
				
				# Obtener o crear el elemento w:tcBorders
				tcPr = celda_borde.get_or_add_tcPr()
				tcBorders = tcPr.first_child_found_in("w:tcBorders")
				if tcBorders is None:
					tcBorders = OxmlElement("w:tcBorders")
					tcPr.append(tcBorders)

				# Agregar los elementos de borde al elemento w:tcBorders
				for borde in ('top', 'left', 'bottom', 'right'):
					borde_elemento = OxmlElement(f'w:{borde}')
					borde_elemento.set(qn('w:val'), 'single')
					borde_elemento.set(qn('w:sz'), '4')
					borde_elemento.set(qn('w:space'), '0')
					borde_elemento.set(qn('w:color'), 'auto')
					tcBorders.append(borde_elemento)

		encabezados = ["NOMBRE DE LA EMPRESA", "NIF DE LA EMPRESA", "EMAIL", "PERSONA DE CONTACTO"]
		for i, fila in enumerate(tabla.rows):
			for j, celda in enumerate(fila.cells):
				if i == 0:
					run = celda.paragraphs[0].add_run(encabezados[j])
				else:
					run = celda.paragraphs[0].add_run(f"tabla{i-1}.{j}")
				
				# Cambiar el tamaño de la fuente
				run.font.size = Pt(10)

		# Buscar el párrafo objetivo
		target_paragraph = None
		for paragraph in documento.paragraphs:
			if contenido_objetivo in paragraph.text:
				target_paragraph = paragraph
				break

		# Insertar la tabla después del párrafo objetivo
		if target_paragraph is not None:
			# Obtener el elemento del párrafo objetivo
			target_element = target_paragraph._element

			# Mover la tabla al lugar deseado
			table_element = tabla._element
			target_element.addnext(table_element)

			# Guardar el documento con la tabla insertada
			documento.save(archivo_salida)
		else:
			print("No se encontró el párrafo objetivo en el documento.")
	except Exception as e:
			error_message = f"Ocurrió un error al insertar la tabla en el documento: {str(e)}"
			QtWidgets.QMessageBox.critical(None, "Error", error_message)
			print(str(e))
def insertar_tabla_en_documento_adjudicacion(archivo_entrada, archivo_salida, contenido_objetivo, filas):

	#try:
		# Abrir el archivo de entrada
		documento = Document(archivo_entrada)

		# Crear una tabla con filas+1 y 4 columnas
		columnas = 4
		tabla = documento.add_table(rows=filas+1, cols=columnas)

		# Establecer el ancho de columna
		ancho_columna = Inches(2)
		for columna in tabla.columns:
			columna.width = round(ancho_columna)

		# Aplicar un borde a todas las celdas de la tabla
		for fila in tabla.rows:
			for celda in fila.cells:
				celda_borde = celda._element
				celda_borde.set(qn('w:borders'), tostring(OxmlElement('w:borders'), encoding='unicode'))
				
				# Obtener o crear el elemento w:tcBorders
				tcPr = celda_borde.get_or_add_tcPr()
				tcBorders = tcPr.first_child_found_in("w:tcBorders")
				if tcBorders is None:
					tcBorders = OxmlElement("w:tcBorders")
					tcPr.append(tcBorders)

				# Agregar los elementos de borde al elemento w:tcBorders
				for borde in ('top', 'left', 'bottom', 'right'):
					borde_elemento = OxmlElement(f'w:{borde}')
					borde_elemento.set(qn('w:val'), 'single')
					borde_elemento.set(qn('w:sz'), '4')
					borde_elemento.set(qn('w:space'), '0')
					borde_elemento.set(qn('w:color'), 'auto')
					tcBorders.append(borde_elemento)

		encabezados = ["Nombre","¿PRESENTA DE OFERTA?","IMPORTE DE LA OFERTA","ORDEN CLASIFICATORIO"]

		orden = [0, 7, 6, 5]

		for i, fila in enumerate(tabla.rows):
			for j, celda in enumerate(fila.cells):
				if i == 0:
					run = celda.paragraphs[0].add_run(encabezados[j])
				else:
					run = celda.paragraphs[0].add_run(f"tabla{i-1}.{orden[j]}")

				# Cambiar el tamaño de la fuente
				run.font.size = Pt(10)
				
				# Cambiar el tamaño de la fuente
				run.font.size = Pt(10)

		# Buscar el párrafo objetivo
		target_paragraph = None
		for paragraph in documento.paragraphs:
			if contenido_objetivo in paragraph.text:
				target_paragraph = paragraph
				break

		# Insertar la tabla después del párrafo objetivo
		if target_paragraph is not None:
			# Obtener el elemento del párrafo objetivo
			target_element = target_paragraph._element

			# Mover la tabla al lugar deseado
			table_element = tabla._element
			target_element.addnext(table_element)

			# Guardar el documento con la tabla insertada
			documento.save(archivo_salida)
		else:
			print("No se encontró el párrafo objetivo en el documento.")
	# except Exception as e:
			# error_message = f"Ocurrió un error al insertar la tabla en el documento: {str(e)}"
			# QtWidgets.QMessageBox.critical(None, "Error", error_message)
			# print(str(e))

